########################################## GeoPlote #######################################

import simplekml
import subprocess
from tkinter import *

from tkinter import ttk
from tkinter import messagebox

import pandas as pd
import math

import utm

import datetime

from tkinter import filedialog
import os

import docx
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import Length


#Listas globais
global data_id
global data_titulos
global data_marcador
global data_obs
global data_coord_ini
global data_coord_fim
global data_zona
global data_qtd
global data_un




#Variáveis para cálculos
#Phi
pi = 3.1415
#he mean radius of Earth is 6,371.0 kilometers and the mean radius of Earth's Moon is 1,737.5 kilometers.
raio_terra = 6371.0


#Função do SimpleKml
kml = simplekml.Kml()


#Lista para coordenadas em Graus Decimais
lista_coord_gd = []

#Lista para coordenadas em Graus Decimais
lista_coord_utm = []


#Valores para controle
i = 0



global pontos
global linhas
# 
# #Conteiner para os pontos gerados
# pontos = kml.newdocument(name = 'Pontos')
# linhas = kml.newdocument(name = 'Linhas')
        
        
        
        
num_itens = 0
contar_linhas = 0



#Listas para a coordenada em UTM início
lista_coord_utm_fuso_ini = []
lista_coord_utm_e_ini = []
lista_coord_utm_s_ini = []
    
#Listas para a coordenada em UTM final
lista_coord_utm_fuso_fim = []
lista_coord_utm_e_fim = []
lista_coord_utm_s_fim = []
    

#Criando listas para as informações do DataFrame
data_id = [] 
    
    
    

#DataFrame dos itens cadastrados
global data_itens
global data_geral

data_itens = pd.DataFrame(columns = ['ID', 'Título',  'Observação',  'Ponto inicial', 'Ponto final', 'Quantidade', 'Unidade'])
print(data_itens)

data_geral = pd.DataFrame(columns = ['ID', 'Título',  'Observação',  'Ponto inicial', 'Ponto final', 'Quantidade', 'Unidade', 'Zona1', 'Cor', 'Marcador', 'largura_linha'])
print(data_geral)

data_geral_temp = pd.DataFrame(columns = ['ID', 'Título',  'Observação',  'Ponto inicial', 'Ponto final', 'Quantidade', 'Unidade', 'Zona1', 'Cor', 'Marcador'])


def verificando_ponto_utm(self, criar_kml):
    print('verificando_ponto_utm')
    #Informações para plotar no mapa. Coordenadas em Graus Decimais
    coord_ini_utm = self.inf_ini_utm.get()
    coord_ini_utm = coord_ini_utm.replace(' ', '')
    
    utm_zona = self.combobox_utm_zonas.get()   
    
    
###################### Adequando informação da coordenada para que o SimpleKml possa ser utilizado #############################


    #Coordenada inicial
    str(coord_ini_utm)

    #Fatiando coordenada inicial fuso
    coord_utm_fuso = coord_ini_utm[0:2]
    #Fatiando coordenada inicial
    coord_utm_e = coord_ini_utm[3:9]
    coord_utm_s = coord_ini_utm[10:17]
    
     
     #Convertendo UTM para graus decimais
    coord_ini_utm = utm.to_latlon(float(coord_utm_e), float(coord_utm_s) , int(coord_utm_fuso), (utm_zona))
    
    print('verificando_ponto_utm_ok')


def verificando_ponto_gd(self, criar_kml):
    print('Verificando_ponto_gd')
    
    #Buscando informações para validá-las    
    coord_ini_gd = self.inf_ini_gd.get()
    coord_ini_gd = coord_ini_gd.replace(' ', '')
    coord_ini_gd = str(coord_ini_gd)

    print('Fatiar_coordenada_GD_ponto')
    coord_ini_gd_lat = coord_ini_gd[0:10]
    coord_ini_gd_lat = (float(coord_ini_gd_lat))
    print(coord_ini_gd_lat)
    coord_ini_gd_lon = coord_ini_gd[11:22]
    coord_ini_gd_lon = (float(coord_ini_gd_lon))

    coord_lista_gd = [coord_ini_gd_lat, coord_ini_gd_lon]
    print(coord_lista_gd[0])
    print(coord_lista_gd[1])
    
    coord_utm = utm.from_latlon(coord_lista_gd[0], coord_lista_gd[1])
    print(coord_utm)
    print('Verificando_ponto_gd_ok')
    
def up_data_delete():
    print('up_data_delete')
    
    global data_geral
    global data_geral_temp
    
    data_geral_lin = data_geral.shape[0]
    
    #Dados por coluna
    id_geral = data_geral['ID']
    titulo_geral = data_geral['Título']
    obs_geral = data_geral['Observação']
    ini_geral = data_geral['Ponto inicial']
    fim_geral = data_geral['Ponto final']
    qtd_geral = data_geral['Quantidade']
    un_geral = data_geral['Unidade']
    zona_geral = data_geral['Zona1']
    cor_geral = data_geral['Cor']
    marc_geral = data_geral['Marcador']
    largura_geral = data_geral['largura_linha']
    
        
    lista_ids = []
    
    for linhas in range(data_geral_lin):
    
        lista_ids.append(linhas)
    
    
    
    data_geral_temp = pd.DataFrame({'ID' : lista_ids, 'Título' : titulo_geral,  'Observação' : obs_geral,  'Ponto inicial' : ini_geral, 'Ponto final' : fim_geral,
                               'Quantidade' : qtd_geral, 'Unidade' : un_geral, 'Zona1' : zona_geral, 'Cor' : cor_geral, 'Marcador' : marc_geral, 'largura_linha' : largura_geral})


    
    data_geral = pd.DataFrame(columns = ['ID', 'Título',  'Observação',  'Ponto inicial', 'Ponto final', 'Quantidade', 'Unidade', 'Zona1', 'Cor', 'Marcador', 'largura_linha'])
    
    data_geral = pd.DataFrame({'ID' : lista_ids, 'Título' : titulo_geral,  'Observação' : obs_geral,  'Ponto inicial' : ini_geral, 'Ponto final' : fim_geral,
                               'Quantidade' : qtd_geral, 'Unidade' : un_geral, 'Zona1' : zona_geral, 'Cor' : cor_geral, 'Marcador' : marc_geral, 'largura_linha' : largura_geral})
            

    
    
def up_data_salvar(self, criar_kml):
    print('up_data_salvar')
    
    global data_itens
    
    
    data_itens = pd.DataFrame(columns = ['ID', 'Título',  'Observação',  'Ponto inicial', 'Ponto final', 'Quantidade', 'Unidade', 'Zona1', 'Cor', 'Marcador'])
#     print(data_itens)
    

    
    #Atualizando data_itens com os dados do self.treeview_cadastros
    for row_id in self.treeview_cadastros.get_children():
        
        #Valores de cada linha do self.treeview_cadastros
        row = self.treeview_cadastros.item(row_id)['values']

        novos_itens = {'ID' : row[0], 'Título' : row[1], 'Observação' : row[2], 'Ponto inicial' : row[3], 'Ponto final' : row[4], 'Quantidade' : row[5], 'Unidade': row[6]}
        #Adicionando valores salvos
        data_itens = data_itens.append(novos_itens, ignore_index=True)                           
        

    
    
    #Atualizando treeview_cadastros com os dados do self. data_itens
    
    #Quantidades de linhas do data_itens
    lin_data_itens = data_itens.shape[0]
#     print(lin_data_itens)
#     print('lin_data_itens')
    
    #Deletando os dados do self.treeview_cadastros
    for i in self.treeview_cadastros.get_children():
        self.treeview_cadastros.delete(i)
#         print(i)
    
    
    
    
    #self.treeview_cadastros aparência
    self.treeview_cadastros.tag_configure('lavender', background = 'lavender', foreground = 'black')
    self.treeview_cadastros.tag_configure('white', background = 'white', foreground = 'black')
    #Adicionando a função de estilo do tkinter
    style = ttk.Style(criar_kml)
    
#     print(data_itens)
    
    #Adicionando itens ao  self.treeview_cadastros
    for lin in range(lin_data_itens):    
         
        #Dados por linha        
        dados_lin = data_itens.loc[[lin]]
#         print(dados_lin)
#         print('dados_lin')
#         print(type(dados_lin))
        
        #Dados por coluna
        id_tree = data_itens['ID']
        titulo_tree = data_itens['Título']
        obs_tree = data_itens['Observação']
        ini_tree = data_itens['Ponto inicial']
        fim_tree = data_itens['Ponto final']
        qtd_tree = data_itens['Quantidade']
        un_tree = data_itens['Unidade']
        
            
        if (lin % 2 != 0):
            style.configure("Treeview", rowheight = 20, font = ('calibri', 11))
#         style.configure("Treeview",
#                 background = "white",
#                 foreground = "black",
#                 rowheight = 25,
#                 fieldbackground = "silver")
        
            self.treeview_cadastros.insert('', 'end', text = lin,
                                   values = (lin, titulo_tree[lin], obs_tree[lin], ini_tree[lin], fim_tree[lin], qtd_tree[lin], un_tree[lin]), tags = ('white'))
        if (lin % 2 == 0):
            style.configure("Treeview", rowheight = 20, font = ('calibri', 11))
            style.configure("Treeview.Headings", rowheight = 25, font = ('calibri', 14))
    #         style.configure("Treeview",
    #         background = "cyan",
    #         foreground = "black",
    #         rowheight = 25,
    #         fieldbackground = "silver")
            self.treeview_cadastros.insert('', 'end', text = lin,
                               values = (lin, titulo_tree[lin], obs_tree[lin], ini_tree[lin], fim_tree[lin], qtd_tree[lin], un_tree[lin]), tags = ('lavender'))




def calculo_dist_utm(self, criar_kml):
    print('calculo_dist_utm')
    global distancia       
    
    #Informações para plotar no mapa. Coordenadas em Graus Decimais
    coord_ini_utm = self.inf_ini_utm.get()
    coord_ini_utm = coord_ini_utm.replace(' ', '')
    coord_fim_utm = self.inf_fim_utm.get()
    coord_fim_utm = coord_fim_utm.replace(' ', '')

    
    utm_zona = self.combobox_utm_zonas.get()

###################### Adequando informação da coordenada para que o SimpleKml possa ser utilizado #############################


    #Coordenada inicial
    str(coord_ini_utm)

    #Fatiando coordenada inicial fuso
    coord_ini_utm_fuso = coord_ini_utm[0:2]
    #Fatiando coordenada inicial
    coord_ini_utm_e = coord_ini_utm[3:9]
    coord_ini_utm_s = coord_ini_utm[10:17]
        

    #Fatiando coordenada final  fuso
    coord_fim_utm_fuso = coord_fim_utm[0:2]
    #Fatiando coordenada final    
    coord_fim_utm_e = coord_fim_utm[3:9]
    coord_fim_utm_s = coord_fim_utm[10:17]


    
    #Listas para a coordenada em UTM início
    lista_coord_utm_fuso_ini.append(coord_ini_utm_fuso)
    lista_coord_utm_e_ini.append(coord_ini_utm_e)
    lista_coord_utm_s_ini.append(coord_ini_utm_s)
     
    

#Convertendo coordenadas de Graus Decimais para UTM
#     str(utm_zona)
#     print(coord_ini_utm_fuso)
#     print(utm_zona)
#     print(coord_ini_utm_e)
#     print(coord_ini_utm_s)
    
    cood_ini_utm = utm.to_latlon(float(coord_ini_utm_e), float(coord_ini_utm_s) , int(coord_ini_utm_fuso), (utm_zona))
#     print(cood_ini_utm)
    
    
    coord_ini_utm_lat = cood_ini_utm[0:1]
    coord_ini_utm_lon = cood_ini_utm[1:2]
    
    print(coord_ini_utm_lat)
    print(coord_ini_utm_lon)
    
#Transformando a tupla gerada pela função .utm() para uma string considerando a latitude
    cood_ini_utm_lat = ' '
    for i in coord_ini_utm_lat:
        cood_ini_utm_lat = cood_ini_utm_lat + str(i)
#         print(cood_ini_utm_lat)
    cood_ini_utm_lat.replace(',', '')

#Transformando a tupla gerada pela função .utm() para uma string considerando a longitude
    cood_ini_utm_lon = ' '
    for i in coord_ini_utm_lon:
        cood_ini_utm_lon = cood_ini_utm_lon + str(i)
#         print(cood_ini_utm_lon)
    cood_ini_utm_lon.replace(' ', '')
    
    
#Fatiando a coordenada após a conversão de UTM para Graus Decimais
    str(cood_ini_utm)
    coord_ini_utm_lat = (cood_ini_utm_lat[0:10])
    coord_ini_utm_lon = (cood_ini_utm_lon[0:10])   
    
    
    

    #Coordenada final
    str(coord_fim_utm)
    
    #Coordenada final
    str(coord_fim_utm)

    #Fatiando coordenada final fuso
    coord_fim_utm_fuso = coord_fim_utm[0:2]
    #Fatiando coordenada final
    coord_fim_utm_e = coord_fim_utm[3:9]
    coord_fim_utm_s = coord_fim_utm[10:17]
        

    #Fatiando coordenada final  fuso
    coord_fim_utm_fuso = coord_fim_utm[0:2]
    #Fatiando coordenada final    
    coord_fim_utm_e = coord_fim_utm[3:9]
    coord_fim_utm_s = coord_fim_utm[10:17]


    
    #Listas para a coordenada em UTM final
    lista_coord_utm_fuso_fim.append(coord_fim_utm_fuso)
    lista_coord_utm_e_fim.append(coord_fim_utm_e)
    lista_coord_utm_s_fim.append(coord_fim_utm_s)
     
     

#Convertendo coordenadas de UTM para  Graus Decimais 
#     str(utm_zona)
#     print(coord_fim_utm_fuso)
#     print(utm_zona)
#     print(coord_fim_utm_e)
#     print(coord_fim_utm_s)

    #Verificando se é um ponto ou uma linha


    if ((coord_ini_utm != '') and (coord_fim_utm != '')):
        cood_fim_utm = utm.to_latlon(float(coord_fim_utm_e), float(coord_fim_utm_s) , int(coord_fim_utm_fuso), (utm_zona))
    #     print(cood_fim_utm)
    #         print(cood_fim_utm)

        
        coord_fim_utm_lat = cood_fim_utm[0:1]
        coord_fim_utm_lon = cood_fim_utm[1:2]
        
    #         print(coord_fim_utm_lat)
    #         print(coord_fim_utm_lon)
        
    #Transformando a tupla gerada pela função .utm() para uma string considerando a latitude
        cood_fim_utm_lat = ' '
        for i in coord_fim_utm_lat:
            cood_fim_utm_lat = cood_fim_utm_lat + str(i)
    #             print(cood_fim_utm_lat)
        cood_fim_utm_lat.replace(',', '')

    #Transformando a tupla gerada pela função .utm() para uma string considerando a longitude
        cood_fim_utm_lon = ' '    
        for i in coord_fim_utm_lon:
            cood_fim_utm_lon = cood_fim_utm_lon + str(i)
            print(cood_fim_utm_lon + ' cood_fim_utm_lon')
        cood_fim_utm_lon.replace(' ', '')

    #Fatiando a coordenada após a conversão de UTM para Graus Decimais
        str(cood_fim_utm)
        coord_fim_utm_lat = (cood_fim_utm_lat[0:10])
        coord_fim_utm_lon = (cood_fim_utm_lon[0:10])
        
    if ((coord_ini_utm != '') and (coord_fim_utm == '')):
        print('É ponto')
        pass

#Verificação para plotar linhas no mapa

    if coord_ini_utm != '' and coord_fim_utm != '':        
        
        lat_inicio = math.radians (float(coord_ini_utm_lat))
        lat_inicio = round(lat_inicio, 15)
        lat_final = math.radians (float(cood_fim_utm_lat))
        lat_final = round(lat_final, 15)
        lon_inicio = math.radians (float(coord_ini_utm_lon))
        lon_inicio = round(lon_inicio, 15)
        lon_final = math.radians (float(cood_fim_utm_lon))
        lon_final = round(lon_final, 15)
        
        #Informação das coordenadas para informar o usuário UTM
        ver_coordenadas = (coord_ini_utm_fuso + '-' + coord_ini_utm_e + ':' + coord_ini_utm_s + ' - ' + coord_fim_utm_fuso + '-' + coord_fim_utm_e + ':' + coord_fim_utm_s)
        
        #Cálculo de distância
        dlon = lon_inicio - lon_final
        dlat = lat_inicio - lat_final
        
        
        ########## https://qastack.com.br/programming/19412462/getting-distance-between-two-points-based-on-latitude-longitude ###########
        a = math.sin(dlat / 2)**2 + math.cos(lat_inicio) * math.cos(lat_final) * math.sin(dlon / 2)**2
        a = round(a, 15) 
        c = 2 * math.atan2(math.sqrt(a), math.sqrt(1 - a))
        c = round(c, 15) 

        distance = raio_terra * c
        distance = round(distance, 3)
        
        #Adeque valor da distância para o padrão BR
        distancia = str(distance)
    

        
def calculo_dist_gd(self, criar_kml):
    print('calculo_dist_gd')
    global distancia 
    
    #Informações para plotar no mapa. Coordenadas em Graus Decimais
    coord_ini_gd = self.inf_ini_gd.get()
    coord_ini_gd = coord_ini_gd.replace(' ', '')
    coord_fim_gd = self.inf_fim_gd.get()
    coord_fim_gd = coord_fim_gd.replace(' ', '')

###################### Adequando informação da coordenada para que o SimpleKml possa ser utilizado #############################

    
#     print(coord_ini_utm_lat)
#     print(coord_ini_utm_lon)
    
#Transformando a tupla gerada pela função .utm() para uma string considerando a latitude
    cood_fim_gd_lat = ' '
    for i in coord_ini_gd:
        cood_fim_gd_lat = cood_fim_gd_lat + str(i)
    cood_fim_gd_lat.replace(',', '')

#Transformando a tupla gerada pela função .utm() para uma string considerando a longitude
    cood_fim_gd_lon = ' '    
    for i in coord_fim_gd:
        cood_fim_gd_lon = cood_fim_gd_lon + str(i)
    cood_fim_gd_lon.replace(' ', '')
    
    #Fatiar as cálcular a distância entre os dois pontos considerando GD
    
    coord_ini_gd_lat = cood_fim_gd_lat[0:10]
    coord_ini_gd_lon = cood_fim_gd_lat[12:22]
    
    cood_fim_gd_lat = cood_fim_gd_lon[0:10]
    cood_fim_gd_lon = cood_fim_gd_lon[12:22]
    
#     print(coord_ini_gd_lat)
#     print(cood_fim_gd_lat)
# 
#     print(coord_ini_gd_lat)
#     print(cood_fim_gd_lat)

    #Verificação para plotar linhas no mapa
    
    lat_inicio = math.radians (float(coord_ini_gd_lat))
    lat_inicio = round(lat_inicio, 15)
    lat_final = math.radians (float(cood_fim_gd_lat))
    lat_final = round(lat_final, 15)
    lon_inicio = math.radians (float(coord_ini_gd_lon))
    lon_inicio = round(lon_inicio, 15)
    lon_final = math.radians (float(cood_fim_gd_lon))
    lon_final = round(lon_final, 15)
    
    
    #Cálculo de distância
    dlon = lon_inicio - lon_final
    dlat = lat_inicio - lat_final
    
    print(lat_inicio)
    
    ########## https://qastack.com.br/programming/19412462/getting-distance-between-two-points-based-on-latitude-longitude ###########
    a = math.sin(dlat / 2)**2 + math.cos(lat_inicio) * math.cos(lat_final) * math.sin(dlon / 2)**2
    a = round(a, 15) 
    c = 2 * math.atan2(math.sqrt(a), math.sqrt(1 - a))
    c = round(c, 15) 

    distance = raio_terra * c
    distance = round(distance, 3)
    
    #Adeque valor da distância para o padrão BR
    distancia = str(distance)       
    


#Funções para os botões coordenadas utm


# Adicionando um valor no ID

    

# Função para o botão 'self.button1'
def salvar_item_f(self, criar_kml):
    print('salvar_item_f')
    
    global data_itens
    global data_geral
    
    global num_itens
    global index
    global iid
    global distancia
    
    global inf_lat
    global inf_lon
    
    #Verificação UTM ou GD
    utm = self.coord_utm.get()
    gd = self.coord_gd.get()
    
        
    c_l = (str(num_itens))      
    
    
    #Buscando informações
    
    #Infomações comuns
    inf_name = self.combobox_titulo.get()
    inf_obs = self.inf_obs.get()
    inf_cor = self.combobox_cor.get()
    inf_marcador = self.combobox_marcador.get()
    inf_zona = self.combobox_utm_zonas.get()
    inf_largura = self.combobox_largura.get()
    
    
    #Infomações de coordenadas
    try:        
        inf_lat = self.inf_ini_utm.get()
        inf_lat = inf_lat.replace(' ', '')
        inf_lon = self.inf_fim_utm.get()
        inf_lon = inf_lon.replace(' ', '')
        print('Salvando_UTM')
    except:
        inf_lat = self.inf_ini_gd.get()
        inf_lat = inf_lat.replace(' ', '')
        inf_lon = self.inf_fim_gd.get()
        inf_lon = inf_lon.replace(' ', '')
        print('Salvando_GD')
    
        #Cálculo de distância para o salvamento dos dados/Verificação de ponto
    
    if (utm == 1) and (inf_zona != '') and (inf_lat != '') and (inf_lon != ''):
        print('Sim_calculo_dist_utm')
        calculo_dist_utm(self, criar_kml)
    else:
        self.label_inf_user.config (text = 'Item inválido, pois não foi indicado nenhuma coordenada.')
        print('Não_calculo_dist_utm')
        
    if (gd == 1) and (inf_zona == '') and (inf_lat != '') and (inf_lon != ''):
        print('Sim_calculo_dist_gd')
        calculo_dist_gd(self, criar_kml)
    else:
        print('Não_calculo_dist_gd')
        self.label_inf_user.config (text = 'Item inválido, pois não foi indicado nenhuma coordenada.')
    
    
#Condições de controle para que as informações sejam salvas
    
    #Se a zona não for informada quando a coordenada UTM estiver como opção
    if (utm == 1)  and (inf_zona == ''):
        print('Favor_informar_a_zona.')
        self.label_inf_user.config (text = 'Favor informar a zona.')

    #Verificando se o item é válido para cadastro, caso não seja informado nenhuma coordenada o item não será salvo
    if (inf_lat == '') and (inf_lon == ''):
        print('Item_inválido_pois_não_foi_indicado_nenhuma_coordenada.')
        self.label_inf_user.config (text = 'Item inválido, pois não foi indicado nenhuma coordenada.')
        
        
    n_itens = len(self.treeview_cadastros.get_children())
    n_cor = n_itens + 1
    

    
    #Adicionando itens ao  self.treeview_cadastros
    
    #Condição para unidade de medida e quantidade
    if (inf_lat != '') and (inf_lon != '') and (inf_lat != inf_lon) and (inf_lat != inf_lat[-1] and inf_lon != inf_lon[-1]):
        qtd = distancia
        qtd = qtd.replace('.', ',')
        un = 'km'
    else:
        qtd = 1
        un = 'un'
        
        
    style = ttk.Style(criar_kml)
    
    #self.treeview_cadastros aparência
    self.treeview_cadastros.tag_configure('lavender', background = 'lavender', foreground = 'black')
    self.treeview_cadastros.tag_configure('white', background = 'white', foreground = 'black')
    
    #Cor do item selecionado
    style.map('Treeview', background = [('selected', 'coral')])
    
#Adicionando itens
        
    #Adicionando valores salvos novos_itens. Considerando a coordenada padrão UTM    
    if ((gd == 0) and (utm == 1) and (inf_zona != '')):        
        
        if (inf_lat != '') and (inf_lon != '') and (inf_lat != inf_lon) and (inf_lat != inf_lat[-1] and inf_lon != inf_lon[-1]):           
            print('Cadastro_linha_UTM')
            
            #Item sem valor para o caso
            inf_marcador = ''

            if (n_cor % 2 != 0):
                style.configure("Treeview", rowheight = 20, font = ('calibri', 11))
    #         style.configure("Treeview",
    #                 background = "white",
    #                 foreground = "black",
    #                 rowheight = 25,
    #                 fieldbackground = "silver")
            
                self.treeview_cadastros.insert('', 'end', text = c_l,
                                       values = (n_itens, inf_name, inf_obs, inf_lat, inf_lon, qtd, un, inf_largura), tags = ('lavender'))
            if (n_cor % 2 == 0):
                style.configure("Treeview", rowheight = 20, font = ('calibri', 11))
    #             style.configure("Treeview",
    #             background = "cyan",
    #             foreground = "black",
    #             rowheight = 25,
    #             fieldbackground = "silver")
                self.treeview_cadastros.insert('', 'end', text = c_l,
                                   values = (n_itens, inf_name, inf_obs, inf_lat, inf_lon, qtd, un, inf_largura), tags = ('white'))
            
            
            #Adicionando valores salvos novos_itens
            novos_itens = {'ID' : n_itens, 'Título' : inf_name, 'Observação' : inf_obs, 'Ponto inicial' : inf_lat, 'Ponto final' : inf_lon, 'Quantidade' : distancia, 'Unidade': 'km'}
            data_itens = data_itens.append(novos_itens, ignore_index=True)
            
            #Adicionando valores salvos novos_itens
            novos_geral = {'ID' : n_itens, 'Título' : inf_name, 'Observação' : inf_obs, 'Ponto inicial' : inf_lat, 'Ponto final' : inf_lon, 'Quantidade' : distancia, 'Unidade': 'km', 'Cor' : inf_cor, 'Marcador' : inf_marcador, 'Zona1' : inf_zona, 'largura_linha' : inf_largura}
            data_geral = data_geral.append(novos_geral, ignore_index=True)

            #Informando ao usuário que o item foi salvo
            self.label_inf_user.config (text = 'Cadastro linha UTM.')
        else:
            print('Cadastro_linha_UTM_não')
            pass
            
            
        if (inf_lat != '') and (inf_lon == '') :
            #Coordenada inicial            
            print('Cadastro_ponto_UTM')
            
            try:
                verificando_ponto_utm(self, criar_kml)
                
                #Item sem valor para o caso
                inf_largura = ''
                #Quantidade igual a 1
                distancia = '1'
                print(inf_marcador)
                #Adicionando valores salvos novos_itens
                novos_itens = {'ID' : n_itens, 'Título' : inf_name, 'Observação' : inf_obs, 'Ponto inicial' : inf_lat, 'Ponto final' : inf_lon, 'Quantidade' : distancia, 'Unidade': 'un'}
                data_itens = data_itens.append(novos_itens, ignore_index=True)
                
                #Adicionando valores salvos novos_itens
                novos_geral = {'ID' : n_itens, 'Título' : inf_name, 'Observação' : inf_obs, 'Ponto inicial' : inf_lat, 'Ponto final' : inf_lon, 'Quantidade' : distancia, 'Unidade': 'un', 'Cor' : inf_cor, 'Marcador' : inf_marcador, 'Zona1' : inf_zona}
                data_geral = data_geral.append(novos_geral, ignore_index=True)
                

                if (n_cor % 2 != 0):
                    self.treeview_cadastros.insert('', 'end', text = c_l,
                                           values = (n_itens, inf_name, inf_obs, inf_lat, inf_lon, qtd, un), tags = ('lavender'))
                if (n_cor % 2 == 0):
                    self.treeview_cadastros.insert('', 'end', text = c_l,
                                       values = (n_itens, inf_name, inf_obs, inf_lat, inf_lon, qtd, un), tags = ('white'))
                    
                #Cor do item selecionado
                style.map('Treeview', background = [('selected', 'coral')])
                
                #Informando ao usuário que o item foi salvo
                self.label_inf_user.config (text = 'Cadastro ponto UTM.')
            except:
                self.label_inf_user.config (text = 'Cadastro não salvo, favor verificar a coordenada informada.')
        else:
            print('Cadastro_ponto_UTM_não')
            pass
#             self.label_inf_user.config (text = 'Cadastro não salvo.')

            
    #Adicionando valores salvos novos_itens. Considerando a coordenada padrão GD    
    if ((gd == 1)  and (utm == 0) and (inf_zona == '')):
        
        if (inf_lat != '') and (inf_lon != '') and (inf_lat != inf_lon) and (inf_lat != inf_lat[-1] and inf_lon != inf_lon[-1]):            
            print('Salvando_item_GD_linha')
            
            #Item sem valor para o caso
            inf_marcador = ''
            
            if (n_cor % 2 != 0):
                style.configure("Treeview", rowheight = 20, font = ('calibri', 11))
    #         style.configure("Treeview",
    #                 background = "white",
    #                 foreground = "black",
    #                 rowheight = 25,
    #                 fieldbackground = "silver")
            
                self.treeview_cadastros.insert('', 'end', text = c_l,
                                       values = (n_itens, inf_name, inf_obs, inf_lat, inf_lon, qtd, un, inf_largura), tags = ('lavender'))
            if (n_cor % 2 == 0):
                style.configure("Treeview", rowheight = 20, font = ('calibri', 11))
    #             style.configure("Treeview",
    #             background = "cyan",
    #             foreground = "black",
    #             rowheight = 25,
    #             fieldbackground = "silver")
                self.treeview_cadastros.insert('', 'end', text = c_l,
                                   values = (n_itens, inf_name, inf_obs, inf_lat, inf_lon, qtd, un, inf_largura), tags = ('white'))
            
            
            #Adicionando valores salvos novos_itens
            novos_itens = {'ID' : n_itens, 'Título' : inf_name, 'Observação' : inf_obs, 'Ponto inicial' : inf_lat, 'Ponto final' : inf_lon, 'Quantidade' : distancia, 'Unidade': 'km'}
            data_itens = data_itens.append(novos_itens, ignore_index=True)
            
            #Adicionando valores salvos novos_itens
            novos_geral = {'ID' : n_itens, 'Título' : inf_name, 'Observação' : inf_obs, 'Ponto inicial' : inf_lat, 'Ponto final' : inf_lon, 'Quantidade' : distancia, 'Unidade': 'km', 'Cor' : inf_cor, 'Marcador' : inf_marcador, 'Zona1' : inf_zona, 'largura_linha' : inf_largura}
            data_geral = data_geral.append(novos_geral, ignore_index=True)
            
            #Informando ao usuário que o item foi salvo
            self.label_inf_user.config (text = 'Cadastro linha graus decimais.')
        else:
            print('Cadastro não salvo_Salvando_item_GD_linha')
            self.label_inf_user.config (text = 'Cadastro não salvo.')
            
        if (inf_lat != '') and (inf_lon == ''):
            print('Salvando_item_GD_ponto')
            
            try:                
                #Verificando informação
                verificando_ponto_gd(self, criar_kml)
                
                #Item sem valor para o caso
                inf_largura = ''
                
                distancia = '1'
                #Adicionando valores salvos novos_itens
                novos_itens = {'ID' : n_itens, 'Título' : inf_name, 'Observação' : inf_obs, 'Ponto inicial' : inf_lat, 'Ponto final' : inf_lon, 'Quantidade' : distancia, 'Unidade': 'un'}
                data_itens = data_itens.append(novos_itens, ignore_index=True)
                
                #Adicionando valores salvos novos_itens
                novos_geral = {'ID' : n_itens, 'Título' : inf_name, 'Observação' : inf_obs, 'Ponto inicial' : inf_lat, 'Ponto final' : inf_lon, 'Quantidade' : distancia, 'Unidade': 'un', 'Cor' : inf_cor, 'Marcador' : inf_marcador, 'Zona1' : inf_zona}
                data_geral = data_geral.append(novos_geral, ignore_index=True)
                

                if (n_cor % 2 != 0):
                    self.treeview_cadastros.insert('', 'end', text = c_l,
                                           values = (n_itens, inf_name, inf_obs, inf_lat, inf_lon, qtd, un), tags = ('lavender'))
                if (n_cor % 2 == 0):
                    self.treeview_cadastros.insert('', 'end', text = c_l,
                                       values = (n_itens, inf_name, inf_obs, inf_lat, inf_lon, qtd, un), tags = ('white'))
                
                print('Ponto')
                            
                up_data_salvar(self, criar_kml)    
                
                #Informando ao usuário que o item foi salvo
                self.label_inf_user.config (text = 'Cadastro ponto graus decimais.')
            except:
                self.label_inf_user.config (text = 'Cadastro não salvo, favor verificar a coordenada informada.')
        else:
            print('Cadastro_não_salvo_Salvando_item_GD_ponto')
#             self.label_inf_user.config (text = 'Cadastro não salvo.')
            pass
            
        #Cor do item selecionado
        style.map('Treeview', background = [('selected', 'coral')])
            
        #Atualizando valores de referência
        data_id.append(num_itens)      
        
            
# Função para o botão 'self.button2'
def exportar_utm_f(self, criar_kml):
    print('exportar_utm_f')    
    exportar_kml_f(self, criar_kml)


# Função para o botão 'self.button_exportar_limpar'
def limpar_kml():
    print('limpar_kml')
    pass
 

    

#Função do botão exportar os dados cadastrados
def exportar_kml_f(self, criar_kml):
    
    global data_geral

    global pontos 
    global linhas
    
    global cood_ini_utm
    global gd_ini

    
    #Nome para o arquivo
    nome = self.nome_arq.get()
    
    ini_exporte = data_geral['Ponto inicial']
    n_ini_exporte = data_geral.shape[0]
    
    #Nome para o arquivo    
    nome = self.nome_arq.get()
    #Diretório para salvar o arquivo
    local_kml = self.dir_salvar.get()
        
    if nome == '':
        print('Favor informar o nome do arquivo.')
        pass
    
    #Resetando o arquivo kml    
    kml = simplekml.Kml(name = 'Pontos')
    kml.document.name = None
    
    kml = simplekml.Kml(name = 'Linhas')
    kml.document.name = None
    
    kml = simplekml.Kml(name = 'pontos')
    kml.document.name = None
    
    kml = simplekml.Kml(name = 'linhas')
    kml.document.name = None
    
    #Conteiner para os pontos gerados
    pontos = kml.newdocument(name = 'Pontos')
    linhas = kml.newdocument(name = 'Linhas')
    

    
    if local_kml != '':        
        local_kml = local_kml
        local_kml = ( (local_kml + '/' + nome + '.kml'))
        
        try:
            kml.save (local_kml)
        except:
            print('Diretório_informado_não_pôde_ser_encontrado.')        
            pass
    else:
        print('Arquivo_não_salvo.')
        self.label_inf_user.config (text = 'Verificar os diretórios informados para a configuração inicial, diretório para salvar o arquivo.kml não encontrado.')
        pass
    
    #Novo ponto abrindo arquivo.kml
    #Diretório para executar o GoogleEarth.
    exe_google_earth = self.combobox_google_earth.get()    
    
    if exe_google_earth != '':        
        #Tentar exportar 
        try:
            #Abrindo o arquivo.kml no GoogleEarth
            subprocess.call([exe_google_earth,  local_kml])
            pass
        except:
            print('Verificar_os_diretórios_informados_para_configuração_inicial')
            pass
    else:
        print('Favor_informar_o_diretório_do_Google_Earth.')
        self.label_inf_user.config (text = 'Verificar os diretórios informados para a configuração inicial, diretório do GoogleEarth.exe não encontrado.')
        pass
    
    n_linhas_data_geral = data_geral.shape[0]
    
    #Verificando antes de exportar
    if (exe_google_earth != '') and (local_kml != ''):
        try:
            #Adicionando itens ao  self.treeview_cadastros
            for item in range(n_linhas_data_geral):    
                 
                #Dados por linha        
                dados_lin = data_geral.loc[[item]]
        #         print(dados_lin)
        #         print('dados_lin')
        #         print(type(dados_lin))
                
                #Dados por coluna
                id_exporte = data_geral['ID']
                titulo_exporte = data_geral['Título']
                obs_exporte = data_geral['Observação']
                ini_exporte = data_geral['Ponto inicial']
                fim_exporte = data_geral['Ponto final']
                qtd_exporte = data_geral['Quantidade']
                un_exporte = data_geral['Unidade']
                zona_exporte = data_geral['Zona1']
                cor_exporte = data_geral['Cor']
                marc_exporte = data_geral['Marcador']
                largura_exporte = data_geral['largura_linha']
                               
                
                #Dados por coluna item a item
                id_exporte_item = (id_exporte.loc[item])
                titulo_exporte_item = (titulo_exporte.loc[item])
                obs_exporte_item = (obs_exporte.loc[item])
                ini_exporte_item =  (ini_exporte.loc[item])
                fim_exporte_item = (fim_exporte.loc[item])  
                qtd_exporte_item = (qtd_exporte.loc[item])
                un_exporte_item = (un_exporte.loc[item])
                zona_exporte_item = (zona_exporte.loc[item])
                cor_exporte_item = (cor_exporte.loc[item])
                marc_exporte_item = (marc_exporte.loc[item])
                largura_exporte_item = (largura_exporte.loc[item])

                #Coordenada UTM
                if (zona_exporte_item != '') and (fim_exporte_item == ''):
                    print('Coordenada_UTM_exporte_ponto.')
                
                    ################################## Erro #########################################
                    #Tentar conversão da coordenada
                    try:
                        print('Coordenada_UTM_convertida_ponto.')
                        
                        #Fatiando coordenada para conversão
                        coord_ini_utm_fuso = ini_exporte_item[0:2]        
                        coord_ini_utm_e = ini_exporte_item[3:9]        
                        coord_ini_utm_s = ini_exporte_item[10:18]     
                        
                        print(coord_ini_utm_fuso)
                        print(coord_ini_utm_e)
                        print(coord_ini_utm_s)
                        
                        cood_ini_gd= utm.to_latlon(float(coord_ini_utm_e), float(coord_ini_utm_s) , int(coord_ini_utm_fuso), (zona_exporte_item))

                        ponto = pontos.newpoint()
                        ponto.coords = [(cood_ini_gd[1], cood_ini_gd[0])]
                        #ponto.name = (titulo_exporte_item)
                        id_exporte_item = str(id_exporte_item)
                        ponto.description = (id_exporte_item + '. ' + titulo_exporte_item + ' _ ' + obs_exporte_item + ' _ ' + ini_exporte_item)
                        
                        #Verificando escolha do marcador. Se o marcador não for escolhido deixar um padrão
                        if marc_exporte_item == '':
                            marc_exporte_item =  'http://maps.google.com/mapfiles/kml/pushpin/wht-pushpin.png'
                            ponto.style.iconstyle.icon.href = (marc_exporte_item)
                        else:
                            ponto.style.iconstyle.icon.href = (marc_exporte_item)
                            
                        #Verificando escolha da cor. Se a cor não for escolhida deixar uma padrão
                        if cor_exporte_item == 'Amarelo':
                            ponto.style.iconstyle.color = simplekml.Color.yellow
                        elif cor_exporte_item == 'Azul':
                            ponto.style.iconstyle.color = simplekml.Color.royalblue
                        elif cor_exporte_item == 'Verde':
                            ponto.style.iconstyle.color = simplekml.Color.limegreen
                        elif cor_exporte_item == 'Vermelho':
                            ponto.style.iconstyle.color = simplekml.Color.orangered
                        else:
                            ponto.style.linestyle.color = simplekml.Color.white
                            
                        ponto.style.iconstyle.scale = 1                        
                    except:
                        print('Coordenada_UTM_não_convertida_ponto.')
                        pass
                    
                if (zona_exporte_item != '') and (ini_exporte_item != '') and (fim_exporte_item != ''):    
                        print('Coordenada_UTM_exporte_linha.')

                        try:
                            print('Coordenada_UTM_convertida_linha.')
                            
                            #Fatiando coordenada para conversão
                            coord_ini_utm_fuso = ini_exporte_item[0:2]        
                            coord_ini_utm_e = ini_exporte_item[3:9]        
                            coord_ini_utm_s = ini_exporte_item[10:18]
                            
                            coord_fim_utm_fuso = fim_exporte_item[0:2]        
                            coord_fim_utm_e = fim_exporte_item[3:9]        
                            coord_fim_utm_s = fim_exporte_item[10:18]
                        
                            cood_ini_utm = utm.to_latlon(float(coord_ini_utm_e), float(coord_ini_utm_s) , int(coord_ini_utm_fuso), (zona_exporte_item))
                            cood_fim_utm = utm.to_latlon(float(coord_fim_utm_e), float(coord_fim_utm_s) , int(coord_fim_utm_fuso), (zona_exporte_item))
                            
                            linha = linhas.newlinestring(coords = [(float(cood_ini_utm[1]), float(cood_ini_utm[0])), (float(cood_fim_utm[1]), float(cood_fim_utm[0]))])

                             #Informações para a linha
                            linha.name = (titulo_exporte_item)
                            qtd_exporte_item = str(qtd_exporte_item)
                            qtd_exporte_item = qtd_exporte_item.replace('.', ',')
                            ver_coordenadas = (coord_ini_utm_fuso + '-' + coord_ini_utm_e + ':' + coord_ini_utm_s + ' _ ' + coord_fim_utm_fuso + '-' + coord_fim_utm_e + ':' + coord_fim_utm_s)
                            id_exporte_item = str(id_exporte_item)
                            linha.description = (id_exporte_item + '. ' + obs_exporte_item + ' _ ' + ver_coordenadas + ' _ ' + str(qtd_exporte_item) + ' km')
                            
                            #Caracteristicas visuais da linha
                            if cor_exporte_item == 'Amarelo':
                                linha.style.linestyle.color = simplekml.Color.yellow
                            elif cor_exporte_item == 'Azul':
                                linha.style.linestyle.color = simplekml.Color.royalblue
                            elif cor_exporte_item == 'Verde':
                                linha.style.linestyle.color = simplekml.Color.limegreen
                            elif cor_exporte_item == 'Vermelho':
                                linha.style.linestyle.color = simplekml.Color.orangered
                            else:
                                linha.style.linestyle.color = simplekml.Color.white
                            
                            #Largura da linha plotada
                            if largura_exporte_item == '':
                                linha.style.linestyle.width = 15
                            else:
                                linha.style.linestyle.width = largura_exporte_item
                                
                        except:
                             print('Coordenada_UTM_não_convertida_linha.')
                             pass
                    
                #Fatiando coordenada para conversão
                if (zona_exporte_item == '') and (fim_exporte_item == ''):
                    print('Coordenada_GD_exporte_ponto.')
                    pass
                
                    try:
                        #Se a coordenada for GD
                        #Fatiando coordenada inicial
                        print('Fatiar_coordenada_GD_ponto')
                        coord_ini_gd_lat = ini_exporte_item[0:10]
                        coord_ini_gd_lat = (float(coord_ini_gd_lat))
                        coord_ini_gd_lon = ini_exporte_item[11:22]
                        coord_ini_gd_lon = (float(coord_ini_gd_lon))
                        cood_ini_gd = [coord_ini_gd_lat, coord_ini_gd_lon]
                        
                        ponto = pontos.newpoint()
                        ponto.coords = [(cood_ini_gd[1], cood_ini_gd[0])]
                        #ponto.name = (titulo_exporte_item)
                        id_exporte_item = str(id_exporte_item)
                        ponto.description = (id_exporte_item + '. ' + titulo_exporte_item + ' _ ' + obs_exporte_item + ' _ ' + ini_exporte_item)            
                        
                        #Verificando escolha do marcador
                        if marc_exporte_item == '':
                            marc_exporte_item =  'http://maps.google.com/mapfiles/kml/pushpin/wht-pushpin.png'
                            ponto.style.iconstyle.icon.href = (marc_exporte_item)
                        else:
                            ponto.style.iconstyle.icon.href = (marc_exporte_item)
                            
                        #Verificando escolha da cor
                        if cor_exporte_item == 'Amarelo':
                            ponto.style.iconstyle.color = simplekml.Color.yellow
                        elif cor_exporte_item == 'Azul':
                            ponto.style.iconstyle.color = simplekml.Color.royalblue
                        elif cor_exporte_item == 'Verde':
                            ponto.style.iconstyle.color = simplekml.Color.limegreen
                        elif cor_exporte_item == 'Vermelho':
                            ponto.style.iconstyle.color = simplekml.Color.orangered
                        else:
                            ponto.style.linestyle.color = simplekml.Color.white
                        
                        ponto.style.iconstyle.scale = 1
                    except:
                        print('Não_fatiar_coordenada_GD_ponto')
                        pass
                    
                if (zona_exporte_item == '') and (ini_exporte_item != '') and (fim_exporte_item != ''):  
                    print('Coordenada_GD_exporte_linha.')
                    pass
                
                    try:
                        print('Fatiar_coordenada_GD_linha')
                        #Fatiando coordenada inicial
                        coord_ini_utm_lat = ini_exporte_item[0:10]
                        coord_ini_utm_lat = (float(coord_ini_utm_lat))
                        coord_ini_utm_lon = ini_exporte_item[11:22]
                        coord_ini_utm_lon = (float(coord_ini_utm_lon))
                        cood_ini_utm = [coord_ini_utm_lat, coord_ini_utm_lon]
                        #Fatiando coordenada final
                        coord_fim_lat = fim_exporte_item[0:10]
                        coord_fim_lat = (float(coord_fim_lat))
                        coord_fim_lon = fim_exporte_item[11:21]
                        coord_fim_lon = (float(coord_fim_lon))                       
                        cood_fim_utm = [coord_fim_lat, coord_fim_lon]
                        
                        #Nova linha GD
                        linha = linhas.newlinestring(coords = [(float(cood_ini_utm[1]), float(cood_ini_utm[0])), (float(cood_fim_utm[1]), float(cood_fim_utm[0]))])
                        
                        #Texto para a legenda da nova linha GD
                        ver_coordenadas = (str(coord_ini_utm_lat) + ', ' + str(coord_ini_utm_lon) + ' - ' + str(coord_fim_lat) + ', ' + str(coord_fim_lon))
                        
                        #Informações para a linha
                        linha.name = (titulo_exporte_item)
                        linha.name = (titulo_exporte_item)
                        qtd_exporte_item = str(qtd_exporte_item)
                        qtd_exporte_item = qtd_exporte_item.replace('.', ',')
                        linha.description = (obs_exporte_item + ' _ ' + ver_coordenadas + ' _ ' + str(qtd_exporte_item) + ' km')

                        #Caracteristicas visuais da linha
                        if cor_exporte_item == 'Amarelo':
                            linha.style.linestyle.color = simplekml.Color.yellow
                        elif cor_exporte_item == 'Azul':
                            linha.style.linestyle.color = simplekml.Color.royalblue
                        elif cor_exporte_item == 'Verde':
                            linha.style.linestyle.color = simplekml.Color.limegreen
                        elif cor_exporte_item == 'Vermelho':
                            linha.style.linestyle.color = simplekml.Color.orangered
                        else:
                            linha.style.linestyle.color = simplekml.Color.white

                        #Largura da linha plotada
                        if largura_exporte_item == '':
                            linha.style.linestyle.width = 15
                        else:
                            linha.style.linestyle.width = largura_exporte_item
                    except:
                        print('Não_fatiar_coordenada_GD_linha')
                        pass

            #Nome para o arquivo    
            nome = self.nome_arq.get()
            
            ##Novo ponto salvando
            local_kml = self.dir_salvar.get()
            
            #Verificando se o diretório para salvamento do arquivo kml foi indicado
            if local_kml != '':        
                local_kml = local_kml.replace("\*", "/")
                local_kml = ( (local_kml + '/' + nome + '.kml'))
            
            #Tentar salvar o arquivo kml
                try:
                    print('Local_para_salvar_kml.')
                    kml.save (local_kml)
                except:
                    print('Local_para_salvar_kml_não_indicado.')
                    pass
            
                #Novo ponto abrindo arquivo.kml
                    
                #Verificando se o diretório para o GoogleEarth foi indicado
                if exe_google_earth != '':        
                    #Tentar exportar 
                    try:
                        #Abrindo o arquivo.kml no GoogleEarth
                        subprocess.call([exe_google_earth,  local_kml])                
                        print('Informado_o_diretório_do_Google_Earth.')
                        #Informando ao usuário que os itens foram exportados
                        kml.save (local_kml)
                        self.label_inf_user.config (text = 'Exportação para o Google Earth concluída.')
                    except:
                        print('Não_informado_algum_diretório.')
                        self.label_inf_user.config (text = 'Exportação para o Google Earth não pôde ser concluída, favor verificar a configuração inicial.')       
                else:
                    print('Favor_informar_o_diretório_do_Google_Earth.')
                    self.label_inf_user.config (text = 'Exportação para o Google Earth não pôde ser concluída, favor verificar a configuração inicial.')   
                    pass
            else:
                self.label_inf_user.config (text = 'Exportação para o Google Earth não pôde ser concluída, favor verificar a configuração inicial.')
        except:
            self.label_inf_user.config (text = 'Exportação para o Google Earth não pôde ser concluída, favor verificar a configuração inicial.')


################################################### Funções para os botões informativos ############################################################

# Função para o botão 'self.button_inf_ini'
def inf_ini_f(self, criar_kml):
    print('inf_ini_f')
    messagebox.showinfo('Informações iniciais', '''Favor informar:
                                                                                                        -O diretório no qual os arquivos gerados serão salvos.
                                                                                                        -O diretório do GoogleEarth, para facilitar é possível indicar o diretório na planilha de configuração coluna 'google_earth'. Exemplo: D:\Google\Google Earth Pro\googleearth.exe
                                                                                                        -Caso queira indique o nome para os arquivos que serão gerados.''')
    
    
# Função para o botão 'self.button_inf_csv'
def inf_dir_csv_f(self, criar_kml):
    print('inf_dir_csv_f')
    messagebox.showinfo('Importar itens para o cadastro', 'Favor informar o diretório do arquivo.csv no qual os itens cadastrados se encontram. Indicar também o nome do arquivo.')
    

# Função para o botão 'self.button_inf_cadastro'
def inf_cadastro_f(self, criar_kml):
    print('inf_cadastro_f')
    messagebox.showinfo('Caracteristicas dos itens no GoogleEarth', '''Caso queira adicionar informações aos itens que serão plotados no GoogleEarth favor preencher os campos deste quadro.
As informações deste quadro definirão as características estéticas e atributos de cada item plotado.''')



# Função para o botão 'self.button_inf_utm'
def inf_utm_f(self, criar_kml):
    messagebox.showinfo('Coordenada UTM', 'Formato: 12-123456:1234567. Favor indicar a letra que designa a zona.')
    

# Função para o botão 'self.button_inf_gd'
def inf_gd_f(self, criar_kml):
    print('inf_gd_f')
    messagebox.showinfo('Coordenada Graus Decimais', 'Formato: 12.123456 (lat), 12.123456 (long)')
    

#################################################### Funções para controle dos itens cadastrados ######################################################

#Definição para a função
def teste1(event):
    print('Teste de um clique!')   
    select_item = self.treeview_cadastros.selection()
    print(select_item)
    
def teste2(event):
    print('Teste de dois cliques!')
    select_item = treeview_teste.selection()
    print(treeview_teste.item(select_item))
    

#Função para deletar algum item
def deletar_item_f(self, criar_kml):
    print('deletar_item_f')
    
    global data_itens
    global data_geral    
    global item_apagar
    
    n_linhas_data_geral = data_geral.shape[0]
    
    if n_linhas_data_geral != 0:
        
        try:
            select_item = self.treeview_cadastros.selection()
        #     print(self.treeview_cadastros.item(select_item))
        #     print(select_item[0])
            item_selecionado = (select_item[0])
            #Apagando item selecionado
            item_selecionado = (item_selecionado[0])
            
            #Deletando do DataFrame data_itens
            item_apagar = self.treeview_cadastros.item(select_item)
            item_apagar = item_apagar['values'][0]
            
            data_geral = data_geral.drop(item_apagar)
            data_geral.reset_index(drop = True, inplace = True)
            
            #Deletando do treeview treeview_cadastros
            self.treeview_cadastros.delete(select_item)
            
        #     print(data_itens)

            up_data_delete()            
               
            up_data_salvar(self, criar_kml)
            
            print('Item_deletado')
            self.label_inf_user.config (text = 'Item deletado.')
            
        except:
            print('Nenhum_item_deletado')
            self.label_inf_user.config (text = 'Nenhum item deletado.')
    else:
            print('Nenhum_item_deletado_para_ser_deletado')
            self.label_inf_user.config (text = 'Nenhum item para ser deletado.')


#Função verificar_limpar_cadastros
def verificar_limpar_cadastros(self, criar_kml):
    print('verificar_limpar_cadastros')
    #Quantidades de linhas do data_itens
    n_lin_data_geral = data_geral.shape[0]
    n_lin_data_itens = data_itens.shape[0]
    
    if (n_lin_data_geral == 0) and (n_lin_data_itens == 0):
            print('Nenhum_item_cadastrado')
            self.label_inf_user.config (text = 'Nenhum item cadastrado.')
    elif (n_lin_data_geral != 0) and (n_lin_data_itens != 0):       
        limpar_cadastros_pergunta = messagebox.askquestion(title = 'Verificar', message = 'Você deseja apagar os itens cadastrados?', icon = 'warning')
        if limpar_cadastros_pergunta == 'yes':
            limpar_cadastros_f(self, criar_kml)
            print('Itens_deletados')
            self.label_inf_user.config (text = 'Itens deletados.')
        else:
            print('Cadastro_não_apagados')
            self.label_inf_user.config (text = 'Cadastro não apagado.')
           
           
#Função do botão button_exportar_limpar
def limpar_cadastros_f(self, criar_kml):
    print('limpar_cadastros_f')
    
    global data_itens
    global data_geral
    global pontos
    global linhas   

    #Infomação para o usuário            
    print(data_itens)
    data_itens = pd.DataFrame(columns = ['ID', 'Título',  'Observação',  'Ponto inicial', 'Ponto final', 'Quantidade', 'Unidade'])
    data_geral = pd.DataFrame(columns = ['ID', 'Título',  'Observação',  'Ponto inicial', 'Ponto final', 'Quantidade', 'Unidade', 'Zona1', 'Cor', 'Marcador', 'largura_linha'])
#         print(data_itens)
    #Deletando os dados do self.treeview_cadastros
    for i in self.treeview_cadastros.get_children():
        self.treeview_cadastros.delete(i)
#         print(i)

    kml = simplekml.Kml(name = 'Pontos')
    kml.document.name = None
    
    kml = simplekml.Kml(name = 'Linhas')
    kml.document.name = None
    
    kml = simplekml.Kml(name = 'pontos')
    kml.document.name = None
    
    kml = simplekml.Kml(name = 'linhas')
    kml.document.name = None
    
    #Conteiner para os pontos gerados
    pontos = kml.newdocument(name = 'Pontos')
    linhas = kml.newdocument(name = 'Linhas')

    #Nome para o arquivo    
    nome = self.nome_arq.get()

    #Local de salvamento dos arquivos gerados
    local_kml = self.dir_salvar.get()

    if local_kml != '':        
        local_kml = local_kml.replace("\*", "/")
        local_kml = ( (local_kml + '/' + nome + '.kml'))            
        try:
            kml.save (local_kml)
        except:
            print('Diretório_para_salvar_o_kml_não_foi_encontrado.')        
            pass
    else:
        print('Arquivo_não_salvo.')
        self.label_inf_user.config (text = 'Verificar os diretórios informados para a configuração inicial, diretório para salvar o arquivo.kml não encontrado.')
        pass
    
    #Novo ponto abrindo arquivo.kml
    #Diretório para executar o GoogleEarth.
    exe_google_earth = self.combobox_google_earth.get()    
    
    if exe_google_earth != '':            
        #Tentar exportar 
        try:
            #Abrindo o arquivo.kml no GoogleEarth
            subprocess.call([exe_google_earth,  local_kml])

        except:
            print('Verificar_os_diretórios_informados_para_configuração_inicial')
            self.label_inf_user.config (text = 'Verificar os diretórios informados para a configuração inicial.')
    else:
        print('Favor_informar_o_diretório_do_Google_Earth.')
        self.label_inf_user.config (text = 'Verificar os diretórios informados para a configuração inicial.')
        pass


#Função do botão  self.button_exportar_csv
def exporte_csv_f(self, criar_kml):
    print('exporte_csv_f')
    global data_geral
    
    #Nome para o arquivo    
    nome = self.nome_arq.get()
    nome = (nome + '.csv')
    
    #Salvar no mesmo diretório
    local_kml = self.dir_salvar.get()
    local_kml_i = local_kml
    local_kml = local_kml.replace("\*", "/")
    
    if local_kml != '':
        try:
            local_kml = (local_kml + '/' + nome)
            data_geral.to_csv(local_kml, index = False, header = True)
            print('Arquivo_csv_salvo')
            self.label_inf_user.config (text = 'Arquivo.csv salvo em ' + local_kml_i + '.')
        except:
            print('Diretório_informado_não_pôde_ser_encontrado.')
            self.label_inf_user.config (text = 'Diretório informado não pôde ser encontrado.')
    elif local_kml == '':
        try:
            local_csv = filedialog.asksaveasfilename(defaultextension = '.csv')
            data_geral.to_csv(local_csv, index = False, header = True)
            print('Salvamento_arquivo.csv_concluído.')
            self.label_inf_user.config (text = 'Exportação do arquivo.csv concluída.')
        except:
            print('Salvamento_arquivo.csv_não_concluído.')
            self.label_inf_user.config (text = 'Exportação do arquivo.csv não concluída.')


#Função do botão  self.button_exportar_csv
def importar_itens_f(self, criar_kml):
    print('importar_csv_f.')
    
    global data_geral
    global data_itens
    global novos_itens
    
    #Quantidades de linhas do data_itens
    n_lin_data_geral = data_geral.shape[0]
    n_lin_data_itens = data_itens.shape[0]


    if (n_lin_data_geral != 0) and (n_lin_data_itens != 0):
        
        try:
            limpar_cadastros_pergunta = messagebox.askquestion(title = 'Verificar', message = 'Para importar os cadastros os itens cadastrados atualmente serão deletados. Deseja continuar?', icon = 'warning')
            if limpar_cadastros_pergunta == 'yes':
                
                limpar_cadastros_f(self, criar_kml)

                data_geral = pd.read_csv(filedialog.askopenfilename(initialdir = os.getcwd()))
                data_geral = data_geral.fillna('')
                #Quantidades de linhas do data_itens
                n_lin = data_geral.shape[0]
                    
                #self.treeview_cadastros aparência
                self.treeview_cadastros.tag_configure('lavender', background = 'lavender', foreground = 'black')
                self.treeview_cadastros.tag_configure('white', background = 'white', foreground = 'black')
                #Adicionando a função de estilo do tkinter
                style = ttk.Style(criar_kml)
                
                print('Percorrer_o_DataFrame_indicado.')
                #Adicionando itens ao  self.treeview_cadastros
                for lin in range(n_lin):
                    
                    #Dados por linha        
                    data_itens = data_geral.loc[[lin]]

                    #Dados por coluna
                    id_tree = data_itens['ID']
                    titulo_tree = data_itens['Título']
                    obs_tree = data_itens['Observação']
                    ini_tree = data_itens['Ponto inicial']
                    fim_tree = data_itens['Ponto final']
                    qtd_tree = data_itens['Quantidade']
                    un_tree = data_itens['Unidade']
                    
                    
                    if (lin % 2 != 0):
                        style.configure("Treeview", rowheight = 20, font = ('calibri', 11))
            #         style.configure("Treeview",
            #                 background = "white",
            #                 foreground = "black",
            #                 rowheight = 25,
            #                 fieldbackground = "silver")
                    
                        self.treeview_cadastros.insert('', 'end', text = lin,
                                               values = (lin, titulo_tree[lin], obs_tree[lin], ini_tree[lin], fim_tree[lin], qtd_tree[lin], un_tree[lin]), tags = ('white'))
                    if (lin % 2 == 0):
                        style.configure("Treeview", rowheight = 20, font = ('calibri', 11))
            #             style.configure("Treeview",
            #             background = "cyan",
            #             foreground = "black",
            #             rowheight = 25,
            #             fieldbackground = "silver")
                        self.treeview_cadastros.insert('', 'end', text = lin,
                                           values = (lin, titulo_tree[lin], obs_tree[lin], ini_tree[lin], fim_tree[lin], qtd_tree[lin], un_tree[lin]), tags = ('lavender'))
                        

                #Cor do item selecionado
                style.map('Treeview', background = [('selected', 'coral')])
                print('Usuário_informou_o_diretório.')
        except:
            print('Usuário não informou o diretório.')
            self.label_inf_user.config (text = 'A ação de importação não foi concluída.')
    else:
        print('Não_existem_itens_para_serem_apagados.')
        
        try:
            print('Tentar_importar_os_itens.')
            data_geral = pd.read_csv(filedialog.askopenfilename(initialdir = os.getcwd()))
            data_geral = data_geral.fillna('')
            #Quantidades de linhas do data_itens
            n_lin = data_geral.shape[0]
                
            #self.treeview_cadastros aparência
            self.treeview_cadastros.tag_configure('lavender', background = 'lavender', foreground = 'black')
            self.treeview_cadastros.tag_configure('white', background = 'white', foreground = 'black')
            #Adicionando a função de estilo do tkinter
            style = ttk.Style(criar_kml)
            
            print('Percorrer_o_DataFrame_indicado.')
            #Adicionando itens ao  self.treeview_cadastros
            for lin in range(n_lin):
                
                #Dados por linha        
                data_itens = data_geral.loc[[lin]]

                #Dados por coluna
                id_tree = data_itens['ID']
                titulo_tree = data_itens['Título']
                obs_tree = data_itens['Observação']
                ini_tree = data_itens['Ponto inicial']
                fim_tree = data_itens['Ponto final']
                qtd_tree = data_itens['Quantidade']
                un_tree = data_itens['Unidade']
                
                #Tratando item quantidade                
                qtd_tree = qtd_tree[lin]
                qtd_tree = str(qtd_tree)
                qtd_tree = qtd_tree.replace('.', ',')

                
                if (lin % 2 != 0):
                    style.configure("Treeview", rowheight = 20, font = ('calibri', 11))
        #         style.configure("Treeview",
        #                 background = "white",
        #                 foreground = "black",
        #                 rowheight = 25,
        #                 fieldbackground = "silver")
                
                    self.treeview_cadastros.insert('', 'end', text = lin,
                                           values = (lin, titulo_tree[lin], obs_tree[lin], ini_tree[lin], fim_tree[lin], qtd_tree, un_tree[lin]), tags = ('white'))
                if (lin % 2 == 0):
                    style.configure("Treeview", rowheight = 20, font = ('calibri', 11))
        #             style.configure("Treeview",
        #             background = "cyan",
        #             foreground = "black",
        #             rowheight = 25,
        #             fieldbackground = "silver")
                    self.treeview_cadastros.insert('', 'end', text = lin,
                                       values = (lin, titulo_tree[lin], obs_tree[lin], ini_tree[lin], fim_tree[lin], qtd_tree, un_tree[lin]), tags = ('lavender'))
                    

            #Cor do item selecionado
            style.map('Treeview', background = [('selected', 'coral')])
            #Informando ao usuário que o item foi salvo
            self.label_inf_user.config (text = 'O arquivo.csv selecionado foi importado.')
        except:
            print('Não foi possível importar os itens.')
            self.label_inf_user.config (text = 'A ação de importação não foi concluída.')
    
    
    
#Função do botão  self.button_exportar_doc
def exporte_doc_f(self, criar_kml):
    print('exporte_doc_f')
    
    #Nome para o arquivo
    nome = self.nome_arq.get()
    
    try:
        local_doc = filedialog.asksaveasfilename(defaultextension = '.docx')
        
        #Criando um documento
        documento = docx.Document()

        #Título do documento
        documento.add_heading(nome, 0)
        
        #Horário em que a exportação do documento foi realizada
        #Buscando a data

        #Data atual para ser indicada no documento exportado
        data_atual = datetime.datetime.now()
        data_inf = data_atual.strftime('%d/%m/%Y às %H:%M horas')
        
        #Valores cadastrados no treeview para serem exportados para o documento

        for linha_id in self.treeview_cadastros.get_children():
            
            #Valores de cada linha do self.treeview_cadastros
            valores_linha = self.treeview_cadastros.item(linha_id)['values']
            
            #Cada valor da linha
            id_linha = valores_linha[0]
            id_linha = str(id_linha)
            
            titulo_linha = valores_linha[1]
            titulo_linha = str(titulo_linha)
            
            obs_linha = valores_linha[2]
            obs_linha = str(obs_linha)
            
            ini_linha = valores_linha[3]
            ini_linha = str(ini_linha)
            
            fim_linha = valores_linha[4]
            fim_linha = str(fim_linha)
            
            qtd_linha = valores_linha[5]
            
            un_linha = valores_linha[6]
            un_linha = str(un_linha)
            
            #Tratamento para os valores de distância
            if un_linha == 'km':
                try:
                    qtd_linha = float(qtd_linha)
                    qtd_linha = (qtd_linha / 1000)
                except:
#                     print('Item_convertido.')
                    pass
                qtd_linha = str(qtd_linha)
                qtd_linha = qtd_linha.replace('.', ',')
            else:
                qtd_linha = str(qtd_linha)
        
                
    #         print(valores_linha)
            
            #Textos para o documento
            if valores_linha[4] == '':
                texto_doc = ('- Item ' + id_linha + '; ' + 'Título: ' + titulo_linha + '; ' + 'Observação: ' + obs_linha + '; ' + 'Coord.ini: ' + ini_linha + '.' )
            elif valores_linha[4] != '':
                texto_doc = ('- Item ' + id_linha + '; ' + 'Título: ' + titulo_linha + '; ' + 'Observação: ' + obs_linha + '; ' +
                                     'Coord.ini: ' + ini_linha + '; ' + 'Coord.fim: ' + fim_linha + '; ' + 'Distância de ' + qtd_linha + ' ' + un_linha + '.')
                
            valores_linha = str(valores_linha)
            linha = documento.add_paragraph().add_run(texto_doc)
            linha.font.size = Pt(12)
            linha.font.name = 'Times New Roman'
        
        #Data do documento
        linha = documento.add_paragraph().add_run('')
        linha = documento.add_paragraph().add_run('Data de referência deste documento ' + data_inf + '.')
        
        #Fim das informações para o docx
        documento.add_page_break()
        
        local_doc = (local_doc)
        documento.save(local_doc)            
        
        print('Salvamento_arquivo_doc_concluído.')
        self.label_inf_user.config (text = 'Exportação do arquivo.docx concluída.')
    
    except:
        print('Salvamento_arquivo_docx_não_concluído.')
        self.label_inf_user.config (text = 'Exportação do arquivo.docx não concluída.')
            

#Função do botão do menu "Arquivo de ajuda"
def menu_arq_ajuda_f(self, criar_kml):
    print('importar_itens_f')       
    try:
        print('Abrir_página_github_ok.')
        os.system("start \"\" https://github.com/Alexsandro1997/kml_plote")
    except:
        print('Abrir_página_github_não.')
        pass
        
